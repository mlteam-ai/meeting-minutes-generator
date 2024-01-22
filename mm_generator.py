from openai import OpenAI
from docx import Document
from pydub import AudioSegment
import os

class MeetingMinutesGenerator:
    def __init__(self) -> None:
        self.client = OpenAI()
        self.output_folder = "output/"

    def transcribe_audio(self, audio_file_path):
        print("Getting the transcription for the audio...")
        # Create a directory to store the output files
        if not os.path.exists(self.output_folder):
            os.mkdir(self.output_folder)

        # OpenAI limitaiton: max input file size is 25MB. So we have to split the audio to minutes and call whisper-1 for each minute
        audio = AudioSegment.from_file(audio_file_path)

        filename, ext = os.path.splitext(audio_file_path)
        ext = ext[1:]

        # PyDub handles time in milliseconds
        one_minute = 60 * 1000

        start = 0
        j = 1
        text = ""
        while start < len(audio):
            end = start + one_minute
            if end > len(audio): 
                end = len(audio)
            next_chunk = audio[start:end]
            chunk_file_name = f"{self.output_folder}{filename}_{j}.{ext}"
            next_chunk.export(chunk_file_name, format=ext)
            with open(chunk_file_name, 'rb') as chunk_file:
                text += self.client.audio.transcriptions.create(model="whisper-1", file=chunk_file).text + " "
            j+=1
            start = end
        print("Transcription is completed.")
        return text
    
    def meeting_minutes(self, transcription):
        abstract_summary = self.__abstract_summary_extraction(transcription)
        key_points = self.__key_points_extraction(transcription)
        action_items = self.__action_item_extraction(transcription)
        sentiment = self.__sentiment_analysis(transcription)
        return {
            'abstract_summary': abstract_summary,
            'key_points': key_points,
            'action_items': action_items,
            'sentiment': sentiment
        }
    
    def __abstract_summary_extraction(self, transcription):
        print("Getting the summary from GPT...")
        response = self.client.chat.completions.create(
            model="gpt-4",
            temperature=0,
            messages=[
                {
                    "role": "system",
                    "content": "You are a highly skilled AI trained in language comprehension and summarization. I would like you to read the following text and summarize it into a concise abstract paragraph. Aim to retain the most important points, providing a coherent and readable summary that could help a person understand the main points of the discussion without needing to read the entire text. Please avoid unnecessary details or tangential points."
                },
                {
                    "role": "user",
                    "content": transcription
                }
            ]
        )
        print("Got the summary from GPT.")
        return response.choices[0].message.content
    

    def __key_points_extraction(self, transcription):
        print("Getting the key points from GPT...")
        response = self.client.chat.completions.create(
            model="gpt-4",
            temperature=0,
            messages=[
                {
                    "role": "system",
                    "content": "You are a proficient AI with a specialty in distilling information into key points. Based on the following text, identify and list the main points that were discussed or brought up. These should be the most important ideas, findings, or topics that are crucial to the essence of the discussion. Your goal is to provide a list that someone could read to quickly understand what was talked about."
                },
                {
                    "role": "user",
                    "content": transcription
                }
            ]
        )
        print("Got the key points from GPT.")
        return response.choices[0].message.content


    def __action_item_extraction(self, transcription):
        print("Getting the action items from GPT...")
        response = self.client.chat.completions.create(
            model="gpt-4",
            temperature=0,
            messages=[
                {
                    "role": "system",
                    "content": "You are an AI expert in analyzing conversations and extracting action items. Please review the text and identify any tasks, assignments, or actions that were agreed upon or mentioned as needing to be done. These could be tasks assigned to specific individuals, or general actions that the group has decided to take. Please list these action items clearly and concisely."
                },
                {
                    "role": "user",
                    "content": transcription
                }
            ]
        )
        print("Got the action items from GPT.")
        return response.choices[0].message.content
    
    def __sentiment_analysis(self, transcription):
        print("Getting the sentiment analysis from GPT...")
        response = self.client.chat.completions.create(
            model="gpt-4",
            temperature=0,
            messages=[
                {
                    "role": "system",
                    "content": "As an AI with expertise in language and emotion analysis, your task is to analyze the sentiment of the following text. Please consider the overall tone of the discussion, the emotion conveyed by the language used, and the context in which words and phrases are used. Indicate whether the sentiment is generally positive, negative, or neutral, and provide brief explanations for your analysis where possible."
                },
                {
                    "role": "user",
                    "content": transcription
                }
            ]
        )
        print("Got the sentiment analysis from GPT.")
        return response.choices[0].message.content    
    
    def save_as_docx(self, minutes, filename):
        print("Saving as .docx file...")
        doc = Document()
        for key, value in minutes.items():
            # Replace underscores with spaces and capitalize each word for the heading
            heading = ' '.join(word.capitalize() for word in key.split('_'))
            doc.add_heading(heading, level=1)
            doc.add_paragraph(value)
            # Add a line break between sections
            doc.add_paragraph()
        doc.save(self.output_folder + filename)
        print("Saved as .docx file")